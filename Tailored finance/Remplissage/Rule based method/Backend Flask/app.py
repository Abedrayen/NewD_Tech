from flask import Flask, request, jsonify, send_file
from docx import Document
import json
import os

app = Flask(__name__)

# Flatten data function
def flatten_data(data, prefix="", max_enfants=4, max_personnes=1, max_produits=2):
    replacements = {}

    if isinstance(data, dict):
        for key, value in data.items():
            new_key = f"{prefix}{key}" if prefix else key
            if isinstance(value, dict):
                replacements.update(flatten_data(value, prefix=new_key + "_"))
            elif isinstance(value, list):
                # Expand lists for `enfants`
                if "enfants" in new_key:
                    for idx in range(max_enfants):
                        if idx < len(value):
                            replacements.update(flatten_data(value[idx], prefix=f"{new_key}{idx}_"))
                        else:
                            replacements[f"{new_key}{idx}_nom"] = "N/A"
                            replacements[f"{new_key}{idx}_date_naissance"] = "N/A"
                            replacements[f"{new_key}{idx}_a_charge"] = "N/A"
                            replacements[f"{new_key}{idx}_remarques"] = "N/A"

                # Expand lists for `personnes_a_charge`
                elif "personnes_a_charge" in new_key:
                    for idx in range(max_personnes):
                        if idx < len(value):
                            replacements.update(flatten_data(value[idx], prefix=f"{new_key}{idx}_"))
                        else:
                            replacements[f"{new_key}{idx}_nom"] = "N/A"
                            replacements[f"{new_key}{idx}_prenom"] = "N/A"
                            replacements[f"{new_key}{idx}_date_naissance"] = "N/A"
                            replacements[f"{new_key}{idx}_remarques"] = "N/A"

                # Expand lists for `produits_epargne`
                elif "produits_epargne" in new_key:
                    for idx in range(max_produits):
                        if idx < len(value):
                            produit = value[idx]
                            replacements.update(flatten_data(produit, prefix=f"{new_key}{idx}_"))
                            # Checkbox logic
                            replacements[f"{new_key}{idx}_est_actuel_oui"] = "☑" if produit["est_actuel"] else "☐"
                            replacements[f"{new_key}{idx}_est_actuel_non"] = "☐" if produit["est_actuel"] else "☑"
                            replacements[f"{new_key}{idx}_passee_oui"] = "☑" if not produit["est_actuel"] else "☐"
                            replacements[f"{new_key}{idx}_passee_non"] = "☐" if not produit["est_actuel"] else "☑"
                            replacements[f"{new_key}{idx}_mode_directe"] = "☑" if produit["mode_gestion"] == "self-managed" else "☐"
                            replacements[f"{new_key}{idx}_mode_conseillee"] = "☐"
                            replacements[f"{new_key}{idx}_mode_sous_mandat"] = "☐"
                        else:
                            replacements[f"{new_key}{idx}_est_actuel_oui"] = "☐"
                            replacements[f"{new_key}{idx}_est_actuel_non"] = "☐"
                            replacements[f"{new_key}{idx}_passee_oui"] = "☐"
                            replacements[f"{new_key}{idx}_passee_non"] = "☐"
                            replacements[f"{new_key}{idx}_mode_directe"] = "☐"
                            replacements[f"{new_key}{idx}_mode_conseillee"] = "☐"
                            replacements[f"{new_key}{idx}_mode_sous_mandat"] = "☐"
                else:
                    for idx, item in enumerate(value):
                        replacements.update(flatten_data(item, prefix=f"{new_key}{idx}_"))
            else:
                replacements[new_key] = str(value) if value is not None else "N/A"

    # Special handling for `connaissance_financieres`
    if "Client" in data:
        niveau = data["Client"].get("connaissance_financieres", "Moyenne")
        replacements["connaissance_financieres_vous_mauvaise"] = "☑" if niveau == "Mauvaise" else "☐"
        replacements["connaissance_financieres_vous_moyenne"] = "☑" if niveau == "Moyenne" else "☐"
        replacements["connaissance_financieres_vous_bonne"] = "☑" if niveau == "Bonne" else "☐"

    if "conjoint" in data:
        niveau_conjoint = data["conjoint"].get("connaissance_financieres", "Moyenne")
        replacements["connaissance_financieres_conjoint_mauvaise"] = "☑" if niveau_conjoint == "Mauvaise" else "☐"
        replacements["connaissance_financieres_conjoint_moyenne"] = "☑" if niveau_conjoint == "Moyenne" else "☐"
        replacements["connaissance_financieres_conjoint_bonne"] = "☑" if niveau_conjoint == "Bonne" else "☐"
    if "capaciteEpargne" in data:
        epargne = data["capaciteEpargne"]
        replacements["compte_courant_vous_montant"] = str(epargne.get("montant", "N/A"))
        replacements["compte_courant_vous_depuis"] = epargne.get("periodicite", "N/A")
        replacements["compte_courant_conjoint_montant"] = str(epargne.get("montant", "N/A"))
        replacements["compte_courant_conjoint_depuis"] = epargne.get("periodicite", "N/A")

    if "IRPP" in data:
        irpp = data["IRPP"]
        replacements["irpp_annee_derniere"] = irpp.get("annee_derniere", "N/A")
        replacements["irpp_moyenne_client"] = irpp.get("moyenne_client", "N/A")
        replacements["irpp_evolutions_previsibles"] = irpp.get("evolutions_previsibles", "N/A")

    if "IFI" in data:
        ifi = data["IFI"]
        replacements["ifi_annee_derniere"] = ifi.get("annee_derniere", "N/A")
        replacements["ifi_moyenne_client"] = ifi.get("moyenne_client", "N/A")
        replacements["ifi_evolutions_previsibles"] = ifi.get("evolutions_previsibles", "N/A")

    if "Autres_Impots_Acquittes" in data:
        autres_impots = ", ".join(data["Autres_Impots_Acquittes"])
        replacements["autres_impots_acquittes"] = autres_impots
        
        
    return replacements

# Replace placeholders in paragraphs
def replace_placeholders_in_paragraph(paragraph, replacements):
    full_text = "".join(run.text for run in paragraph.runs)
    for key, value in replacements.items():
        placeholder = f"{{{{{key}}}}}"
        full_text = full_text.replace(placeholder, value)
    if full_text != paragraph.text:
        for run in paragraph.runs:
            run.text = ""
        paragraph.add_run(full_text)

# Replace placeholders in tables
def replace_placeholders_in_table(table, replacements):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_placeholders_in_paragraph(paragraph, replacements)

# Process Word Document
def process_document(template_path, output_path, replacements):
    doc = Document(template_path)
    for paragraph in doc.paragraphs:
        replace_placeholders_in_paragraph(paragraph, replacements)
    for table in doc.tables:
        replace_placeholders_in_table(table, replacements)
    doc.save(output_path)
    print(f"Document saved: {output_path}")

@app.route('/process-document', methods=['POST'])
def process_document_endpoint():
    try:
        data = request.json
        if not data:
            return jsonify({"error": "No data provided"}), 400

        # Flatten JSON data
        flattened_data = flatten_data(data)

        # Paths
        template_path = "Document Connaissance Client.docx"
        output_path = "Document_Connaissance_Client_Filled.docx"

        # Process the document
        process_document(template_path, output_path, flattened_data)

        # Return the document
        return send_file(output_path, as_attachment=True)
    except Exception as e:
        return jsonify({"error": str(e)}), 500

if __name__ == '__main__':
    app.run(port=5000)
